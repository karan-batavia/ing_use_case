import os
import sys
import argparse
import glob
from pathlib import Path
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


class DOCXToTextConverter:
    def __init__(
        self,
        include_headers=True,
        include_footers=True,
        include_tables=True,
        preserve_formatting=False,
        clean_whitespace=True,
        remove_empty_lines=True,
        table_separator=" | ",
        include_metadata=False,
    ):
        """
        Initialize the DOCX to text converter with various options.

        Args:
            include_headers (bool): Include header text in output
            include_footers (bool): Include footer text in output
            include_tables (bool): Include table content in output
            preserve_formatting (bool): Attempt to preserve some text formatting
            clean_whitespace (bool): Remove extra whitespace and normalize spacing
            remove_empty_lines (bool): Remove empty lines from output
            table_separator (str): Separator for table cells
            include_metadata (bool): Include document metadata at the beginning
        """
        self.include_headers = include_headers
        self.include_footers = include_footers
        self.include_tables = include_tables
        self.preserve_formatting = preserve_formatting
        self.clean_whitespace = clean_whitespace
        self.remove_empty_lines = remove_empty_lines
        self.table_separator = table_separator
        self.include_metadata = include_metadata

    def extract_metadata(self, document):
        """Extract document metadata."""
        metadata = []
        core_props = document.core_properties

        if core_props.title:
            metadata.append(f"Title: {core_props.title}")
        if core_props.author:
            metadata.append(f"Author: {core_props.author}")
        if core_props.subject:
            metadata.append(f"Subject: {core_props.subject}")
        if core_props.created:
            metadata.append(f"Created: {core_props.created}")
        if core_props.modified:
            metadata.append(f"Modified: {core_props.modified}")
        if core_props.last_modified_by:
            metadata.append(f"Last Modified By: {core_props.last_modified_by}")

        if metadata:
            return (
                "=== DOCUMENT METADATA ===\n"
                + "\n".join(metadata)
                + "\n"
                + "=" * 30
                + "\n\n"
            )
        return ""

    def extract_table_text(self, table):
        """Extract text from a table."""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + " "
                row_text.append(cell_text.strip())
            table_text.append(self.table_separator.join(row_text))
        return "\n".join(table_text)

    def extract_headers_footers(self, document):
        """Extract text from headers and footers."""
        headers_footers_text = []

        if self.include_headers:
            for section in document.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            headers_footers_text.append(f"[HEADER] {paragraph.text}")

        if self.include_footers:
            for section in document.sections:
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            headers_footers_text.append(f"[FOOTER] {paragraph.text}")

        return headers_footers_text

    def process_paragraph(self, paragraph):
        """Process a paragraph with formatting options."""
        if not paragraph.text.strip():
            return ""

        text = paragraph.text

        if self.preserve_formatting:
            # Add basic formatting indicators
            if paragraph.style.name.startswith("Heading"):
                level = (
                    paragraph.style.name[-1]
                    if paragraph.style.name[-1].isdigit()
                    else "1"
                )
                text = f"{'#' * int(level)} {text}"
            elif paragraph.style.name == "Title":
                text = f"# {text}"
            elif paragraph.style.name in ["Quote", "Intense Quote"]:
                text = f"> {text}"

        return text

    def convert_docx_to_txt(self, docx_file, txt_file=None):
        """
        Convert a single DOCX file to text.

        Args:
            docx_file (str): Path to the input DOCX file
            txt_file (str, optional): Path to the output TXT file.
                                     If None, auto-generates based on input filename.

        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            # Auto-generate output filename if not provided
            if txt_file is None:
                input_path = Path(docx_file)
                txt_file = input_path.with_suffix(".txt")

            # Load the DOCX file
            document = Document(docx_file)

            # Initialize text collection
            all_text = []

            # Add metadata if requested
            if self.include_metadata:
                metadata = self.extract_metadata(document)
                if metadata:
                    all_text.append(metadata)

            # Extract headers and footers
            headers_footers = self.extract_headers_footers(document)
            if headers_footers:
                all_text.extend(headers_footers)
                all_text.append("")  # Add separator

            # Process document body
            for element in document.element.body:
                if isinstance(element, CT_P):
                    # It's a paragraph
                    paragraph = Paragraph(element, document)
                    processed_text = self.process_paragraph(paragraph)
                    if processed_text:
                        all_text.append(processed_text)

                elif isinstance(element, CT_Tbl) and self.include_tables:
                    # It's a table
                    table = Table(element, document)
                    table_text = self.extract_table_text(table)
                    if table_text:
                        all_text.append("\n[TABLE]")
                        all_text.append(table_text)
                        all_text.append("[/TABLE]\n")

            # Join all text
            full_text = "\n".join(all_text)

            # Apply text processing options
            if self.clean_whitespace:
                # Normalize whitespace
                lines = full_text.split("\n")
                cleaned_lines = []
                for line in lines:
                    cleaned_line = " ".join(line.split())  # Remove extra whitespace
                    cleaned_lines.append(cleaned_line)
                full_text = "\n".join(cleaned_lines)

            if self.remove_empty_lines:
                # Remove empty lines
                lines = full_text.split("\n")
                non_empty_lines = [line for line in lines if line.strip()]
                full_text = "\n".join(non_empty_lines)

            # Write the extracted text to a TXT file
            with open(txt_file, "w", encoding="utf-8") as file:
                file.write(full_text)

            print(f"✓ Successfully converted '{docx_file}' to '{txt_file}'.")
            return True

        except FileNotFoundError:
            print(f"✗ Error: The file '{docx_file}' was not found.")
            return False
        except Exception as e:
            print(f"✗ An error occurred while converting '{docx_file}': {e}")
            return False

    def convert_multiple_files(self, input_pattern, output_dir=None):
        """
        Convert multiple DOCX files matching a pattern.

        Args:
            input_pattern (str): Glob pattern for input DOCX files
            output_dir (str, optional): Directory for output files. If None, uses same directory as input.

        Returns:
            int: Number of files successfully converted
        """
        docx_files = glob.glob(input_pattern)
        if not docx_files:
            print(f"No DOCX files found matching pattern: {input_pattern}")
            return 0

        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        successful_conversions = 0
        total_files = len(docx_files)

        print(f"Found {total_files} DOCX file(s) to convert...")

        for i, docx_file in enumerate(docx_files, 1):
            print(f"[{i}/{total_files}] Processing: {docx_file}")

            if output_dir:
                output_file = os.path.join(output_dir, Path(docx_file).stem + ".txt")
            else:
                output_file = None  # Auto-generate in same directory

            if self.convert_docx_to_txt(docx_file, output_file):
                successful_conversions += 1

        print(
            f"\nConversion complete! {successful_conversions}/{total_files} files converted successfully."
        )
        return successful_conversions


def interactive_mode():
    """Run the converter in interactive mode."""
    print("=== DOCX to Text Converter - Interactive Mode ===")

    while True:
        print("\nOptions:")
        print("1. Convert single file")
        print("2. Convert multiple files (batch)")
        print("3. Exit")

        choice = input("\nSelect an option (1-3): ").strip()

        if choice == "1":
            input_file = input("Enter DOCX file path: ").strip()
            if not os.path.exists(input_file):
                print(f"File not found: {input_file}")
                continue

            output_file = input(
                "Enter output file path (press Enter for auto-generation): "
            ).strip()
            output_file = output_file if output_file else None

            # Get processing options
            print("\n--- Processing Options ---")
            include_headers = input(
                "Include headers? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            include_footers = input(
                "Include footers? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            include_tables = input(
                "Include tables? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            preserve_formatting = input(
                "Preserve formatting (headings, quotes)? (y/n, default=n): "
            ).strip().lower() in ["y", "yes"]
            include_metadata = input(
                "Include document metadata? (y/n, default=n): "
            ).strip().lower() in ["y", "yes"]
            clean_ws = input(
                "Clean whitespace? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            remove_empty = input(
                "Remove empty lines? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]

            converter = DOCXToTextConverter(
                include_headers=include_headers,
                include_footers=include_footers,
                include_tables=include_tables,
                preserve_formatting=preserve_formatting,
                include_metadata=include_metadata,
                clean_whitespace=clean_ws,
                remove_empty_lines=remove_empty,
            )
            converter.convert_docx_to_txt(input_file, output_file)

        elif choice == "2":
            pattern = input("Enter file pattern (e.g., *.docx, docs/*.docx): ").strip()
            output_dir = input(
                "Enter output directory (press Enter for same directory): "
            ).strip()
            output_dir = output_dir if output_dir else None

            # Get processing options
            print("\n--- Processing Options ---")
            include_headers = input(
                "Include headers? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            include_footers = input(
                "Include footers? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            include_tables = input(
                "Include tables? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            preserve_formatting = input(
                "Preserve formatting? (y/n, default=n): "
            ).strip().lower() in ["y", "yes"]
            include_metadata = input(
                "Include metadata? (y/n, default=n): "
            ).strip().lower() in ["y", "yes"]
            clean_ws = input(
                "Clean whitespace? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]
            remove_empty = input(
                "Remove empty lines? (y/n, default=y): "
            ).strip().lower() in ["", "y", "yes"]

            converter = DOCXToTextConverter(
                include_headers=include_headers,
                include_footers=include_footers,
                include_tables=include_tables,
                preserve_formatting=preserve_formatting,
                include_metadata=include_metadata,
                clean_whitespace=clean_ws,
                remove_empty_lines=remove_empty,
            )
            converter.convert_multiple_files(pattern, output_dir)

        elif choice == "3":
            print("Goodbye!")
            break
        else:
            print("Invalid choice. Please select 1, 2, or 3.")


def main():
    """Main function with command-line argument support."""
    parser = argparse.ArgumentParser(
        description="Convert DOCX files to plain text with advanced options",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python docx_to_txt.py file.docx                     # Convert single file
  python docx_to_txt.py file.docx -o output.txt       # Specify output file
  python docx_to_txt.py "*.docx" --batch              # Convert all DOCX files
  python docx_to_txt.py "docs/*.docx" --batch -d txt  # Batch convert to 'txt' directory
  python docx_to_txt.py --interactive                 # Run in interactive mode
  python docx_to_txt.py file.docx --preserve-formatting --include-metadata
        """,
    )

    parser.add_argument("input", nargs="?", help="Input DOCX file or pattern")
    parser.add_argument(
        "-o", "--output", help="Output text file (for single file conversion)"
    )
    parser.add_argument(
        "-d", "--output-dir", help="Output directory (for batch conversion)"
    )
    parser.add_argument(
        "--batch", action="store_true", help="Batch mode: treat input as file pattern"
    )
    parser.add_argument(
        "--interactive", action="store_true", help="Run in interactive mode"
    )

    # Content options
    parser.add_argument(
        "--no-headers", action="store_true", help="Exclude headers from output"
    )
    parser.add_argument(
        "--no-footers", action="store_true", help="Exclude footers from output"
    )
    parser.add_argument(
        "--no-tables", action="store_true", help="Exclude tables from output"
    )
    parser.add_argument(
        "--preserve-formatting",
        action="store_true",
        help="Preserve text formatting (headings, quotes)",
    )
    parser.add_argument(
        "--include-metadata", action="store_true", help="Include document metadata"
    )

    # Text processing options
    parser.add_argument(
        "--no-clean", action="store_true", help="Don't clean whitespace"
    )
    parser.add_argument(
        "--keep-empty-lines", action="store_true", help="Keep empty lines"
    )
    parser.add_argument(
        "--table-separator",
        default=" | ",
        help='Separator for table cells (default: " | ")',
    )

    args = parser.parse_args()

    # Interactive mode
    if args.interactive:
        interactive_mode()
        return

    # Check if input is provided
    if not args.input:
        print(
            "No input specified. Use --interactive for interactive mode or provide an input file/pattern."
        )
        parser.print_help()
        return

    # Create converter with options
    converter = DOCXToTextConverter(
        include_headers=not args.no_headers,
        include_footers=not args.no_footers,
        include_tables=not args.no_tables,
        preserve_formatting=args.preserve_formatting,
        include_metadata=args.include_metadata,
        clean_whitespace=not args.no_clean,
        remove_empty_lines=not args.keep_empty_lines,
        table_separator=args.table_separator,
    )

    # Batch mode
    if args.batch:
        converter.convert_multiple_files(args.input, args.output_dir)
    else:
        # Single file mode
        if not os.path.exists(args.input):
            print(f"Error: File '{args.input}' not found.")
            sys.exit(1)

        success = converter.convert_docx_to_txt(args.input, args.output)
        sys.exit(0 if success else 1)


# Example usage and backward compatibility
if __name__ == "__main__":
    if len(sys.argv) == 1:
        # No arguments provided, run with default behavior for backward compatibility
        print("Running in legacy mode...")
        converter = DOCXToTextConverter()

        # Try to find a valid DOCX file to demonstrate the converter
        test_files = [
            "tests/data/prompts_01_06_14.docx",
            "tests/data/prompts_04_13_16.docx",
            "tests/data/prompts_07_08_12.docx",
        ]

        input_docx = None
        for test_file in test_files:
            if os.path.exists(test_file):
                input_docx = test_file
                break

        if input_docx is None:
            # Try from script directory
            script_dir = os.path.dirname(os.path.abspath(__file__))
            workspace_root = os.path.dirname(os.path.dirname(script_dir))
            for test_file in test_files:
                full_path = os.path.join(workspace_root, test_file)
                if os.path.exists(full_path):
                    input_docx = full_path
                    break

        if input_docx:
            output_txt = "sample_output.txt"  # Replace with your desired TXT file path
            converter.convert_docx_to_txt(input_docx, output_txt)
        else:
            print("No valid DOCX test files found. Please:")
            print("1. Use command line arguments: python docx_to_txt.py your_file.docx")
            print("2. Or run: python docx_to_txt.py --interactive")
            print("3. Or run: python create_test_docx.py to generate test files")
    else:
        # Arguments provided, use new dynamic interface
        main()
