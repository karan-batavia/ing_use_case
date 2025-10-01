from docx import Document
from docx.shared import Inches
import io
from typing import Optional


class DOCXWriter:
    """Class to write text content to DOCX format"""

    def __init__(self):
        pass

    def create_docx_from_text(
        self, text_content: str, output_path: Optional[str] = None
    ) -> bytes:
        """
        Create a DOCX document from text content

        Args:
            text_content: The text to write to the document
            output_path: Optional path to save the file

        Returns:
            bytes: The DOCX file as bytes
        """
        try:
            # Create a new document
            doc = Document()

            # Add a title
            title = doc.add_heading("Scrubbed Document", 0)

            # Split text into paragraphs and add them
            paragraphs = text_content.split("\n\n")
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Split by single newlines to preserve line breaks
                    lines = paragraph_text.split("\n")
                    for i, line in enumerate(lines):
                        if line.strip():
                            if i == 0:
                                # First line of paragraph
                                p = doc.add_paragraph(line.strip())
                            else:
                                # Add line break and continue text
                                doc.add_paragraph(line.strip())

            # Save to bytes
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Optionally save to file
            if output_path:
                doc.save(output_path)

            return file_stream.getvalue()

        except Exception as e:
            raise Exception(f"Error creating DOCX file: {str(e)}")

    def create_docx_buffer(self, text_content: str) -> io.BytesIO:
        """
        Create a DOCX document and return as BytesIO buffer

        Args:
            text_content: The text to write to the document

        Returns:
            io.BytesIO: Buffer containing the DOCX file
        """
        try:
            doc = Document()

            # Add title
            doc.add_heading("Scrubbed Document", 0)

            # Add content with proper paragraph handling
            paragraphs = text_content.split("\n\n")
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            # Create buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer

        except Exception as e:
            raise Exception(f"Error creating DOCX buffer: {str(e)}")
