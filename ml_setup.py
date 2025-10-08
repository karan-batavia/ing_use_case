#!/usr/bin/env python3
"""
Hybrid Sensitivity Classifier (rules + ML + optional Presidio) with structure-aware multi-format support.

Primary use-case: Train/evaluate on a CSV file with:
  - a text column (e.g. "text")
  - an optional label column (values: C1, C2, C3, C4)

If no label column is provided, the script will auto-label with heuristics.

Run:
  python ml_setup.py --train --data /path/to/data.csv --text-col text --label-col label

Env vars (optional):
  DATA_PATH=/path/to/data.csv
  MODEL=logreg|rf            (default logreg)
  CSV_TEXT_COLUMN=text
  CSV_LABEL_COLUMN=label
  TEST_SIZE=0.2              (float)
  MIN_DF=1                   (int)
  NGRAM_MAX=2                (int)
  USE_PRESIDIO=0|1           (default 1, will skip if Presidio not installed)
"""

import os
import re
import sys
import json
from dataclasses import dataclass
from typing import List, Tuple, Dict, Any, Optional
from pathlib import Path

from sklearn.base import BaseEstimator, TransformerMixin
from sklearn.pipeline import FeatureUnion, Pipeline
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report, confusion_matrix
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.feature_extraction.text import TfidfVectorizer, HashingVectorizer
import joblib
import unicodedata

# --------------------------
# Simple normalization
# --------------------------
def simple_normalize(s: str) -> str:
    # Keep accents; just normalize unicode/spacing
    return unicodedata.normalize("NFKC", str(s)).strip()

# --------------------------
# Optional: Presidio analyzer (skipped if not available)
# --------------------------
USE_PRESIDIO = os.environ.get("USE_PRESIDIO", "1") == "1"
HAS_PRESIDIO = False
try:
    if USE_PRESIDIO:
        try:
            # Prefer your custom analyzer if present
            from src.presidio_pipeline import analyzer as _prebuilt_analyzer  # noqa
            _ANALYZER = _prebuilt_analyzer
            HAS_PRESIDIO = True
            print("Presidio: using analyzer from src.presidio_pipeline")
        except Exception:
            from presidio_analyzer import AnalyzerEngine, RecognizerRegistry  # type: ignore
            _registry = RecognizerRegistry()
            _registry.load_predefined_recognizers()
            _ANALYZER = AnalyzerEngine(registry=_registry, supported_languages=["en"])
            HAS_PRESIDIO = True
            print("Presidio: using default AnalyzerEngine")
except Exception as _e:
    print(f"Presidio unavailable: {_e}")
    HAS_PRESIDIO = False

# --------------------------
# Optional structured extractors
# --------------------------
HAS_PANDAS = False
try:
    import pandas as pd
    HAS_PANDAS = True
except Exception:
    print("Warning: pandas not installed. CSV support disabled.")

HAS_PDFPLUMBER = False
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except Exception:
    pass

HAS_PYPDF2 = False
try:
    import PyPDF2
    HAS_PYPDF2 = True
except Exception:
    pass

HAS_DOCX2PY = False
try:
    from docx2python import docx2python
    HAS_DOCX2PY = True
except Exception:
    pass

HAS_DOCX = False
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    pass

HAS_OCR = False
try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except Exception:
    pass

# --------------------------
# Config
# --------------------------
DEFAULT_DATA = "raw_prompts.txt"
DATA_PATH = os.environ.get("DATA_PATH", DEFAULT_DATA)
MODEL_CHOICE = os.environ.get("MODEL", "logreg").lower()     # logreg | rf
RANDOM_STATE = int(os.environ.get("SEED", "42"))
TEST_SIZE = float(os.environ.get("TEST_SIZE", "0.2"))
NGRAM_MAX = int(os.environ.get("NGRAM_MAX", "2"))
MIN_DF = int(os.environ.get("MIN_DF", "1"))
MAX_DF = float(os.environ.get("MAX_DF", "1.0"))

CSV_TEXT_COLUMN = os.environ.get("CSV_TEXT_COLUMN")  # can be overridden by --text-col
CSV_LABEL_COLUMN = os.environ.get("CSV_LABEL_COLUMN")  # can be overridden by --label-col

# --------------------------
# Patterns (for rules)
# --------------------------
FALLBACK_PATTERNS: Dict[str, str] = {
    "EMAIL": r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
    "PHONE_EU": r"(?:\+\d{1,3}\s?)?(?:\d[\s-]?){9,}",
    "SSN_LIKE": r"\b\d{6}[- ]?\d{2,4}[\.]?\d{0,2}\b",
    "IBAN": r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b",
    "ACCOUNT_NUM": r"\b(?:acct|account)\s*\d{3,}\b",
    "AMOUNT": r"(?:USD|EUR|GBP|€|\$|£)\s?\d{1,3}(?:[, \u00A0]\d{3})*(?:\.\d{2})?",
    "DOB": r"\b\d{4}-\d{2}-\d{2}\b",
    "NATIONAL_ID": r"\bID[:\s-]?[A-Z0-9]{6,}\b",
    "BIOMETRIC": r"\b(FaceID|fingerprint|iris|biometric)\b",
}
def _compile_pat(v) -> re.Pattern:
    if isinstance(v, re.Pattern):
        return re.compile(v.pattern, re.IGNORECASE)
    return re.compile(str(v), re.IGNORECASE)
PATTERNS: Dict[str, re.Pattern] = {k: _compile_pat(v) for k, v in FALLBACK_PATTERNS.items()}

# --------------------------
# Labels + heuristics
# --------------------------
LABELS = ["C1", "C2", "C3", "C4"]

# Accept common label variants and normalize to C1–C4
LABEL_ALIASES = {
    # words
    "LOW": "C1",
    "MEDIUM": "C2",
    "MID": "C2",
    "HIGH": "C3",
    "CRITICAL": "C4",
    "VERY HIGH": "C4",
    # numbers (string)
    "1": "C1", "2": "C2", "3": "C3", "4": "C4",
    # with prefixes/spacing
    "C 1": "C1", "C 2": "C2", "C 3": "C3", "C 4": "C4",
    "C-1": "C1", "C-2": "C2", "C-3": "C3", "C-4": "C4",
}

def normalize_label(val: Any) -> Optional[str]:
    if val is None:
        return None
    s = str(val).strip().upper()
    if s in LABELS:
        return s
    return LABEL_ALIASES.get(s)

def heuristic_label(text: str) -> str:
    t = text.lower()

    # (Optional) Presidio signals
    if USE_PRESIDIO and HAS_PRESIDIO:
        try:
            ents = _ANALYZER.analyze(text=text, language="en")
            types = {r.entity_type for r in ents}
            # Adjust these to your recognizer set if you use a custom registry
            if {"IBAN", "CREDIT_CARD", "SSN", "NATIONAL_ID", "CRYPTO", "BANK_ACCOUNT"}.intersection(types):
                return "C4"
            # crude combo signal:
            if "PERSON" in types and (PATTERNS["EMAIL"].search(text) and PATTERNS["AMOUNT"].search(text)):
                return "C4"
        except Exception:
            pass

    c4_hits = any(PATTERNS[k].search(text) for k in ["SSN_LIKE", "IBAN", "NATIONAL_ID", "BIOMETRIC"])
    email_hit = PATTERNS["EMAIL"].search(text) is not None
    amount_hit = PATTERNS["AMOUNT"].search(text) is not None
    if c4_hits or (email_hit and amount_hit):
        return "C4"

    c3_words = any(w in t for w in [
        "agreement","supplier","customer","standing order","payment order","deposit","overdraft",
        "line of credit","bank guarantee","letter of credit","invoice","po-"
    ])
    if c3_words or amount_hit:
        return "C3"

    c2_words = any(w in t for w in [
        "policy","guideline","standard","sop","governance","raci","owner","approver",
        "review cycle","deprecated","retired","under review","internal","digest","reminder","notice"
    ])
    if c2_words:
        return "C2"

    c1_words = any(w in t for w in [
        "annual report","pillar 3","press","newsroom","full year results","half year results",
        "public disclosures","investor","analyst","linkedin","press-room","pdf","xlsx","report"
    ])
    if c1_words:
        return "C1"

    return "C2"

# --------------------------
# Structure-aware parsing
# --------------------------
def _clean_units(units: List[str]) -> List[str]:
    out = []
    for u in units:
        u = simple_normalize(u)
        if not u:
            continue
        if len(u) < 4 and not re.search(r"[0-9@€$£]", u):
            continue
        out.append(u)
    return out

def parse_txt(path: Path) -> List[str]:
    with path.open("r", encoding="utf-8") as f:
        units = [line.rstrip("\n") for line in f if line.strip()]
    return _clean_units(units)

def parse_csv(path: Path, text_col: Optional[str]) -> Tuple[List[str], Optional[List[str]]]:
    if not HAS_PANDAS:
        raise RuntimeError("pandas required for CSV")
    df = pd.read_csv(path)
    if text_col and text_col not in df.columns:
        raise ValueError(f"CSV text column '{text_col}' not in columns: {list(df.columns)}")

    # Prefer explicit text column; else concat the row
    if text_col:
        texts = df[text_col].astype(str).tolist()
    else:
        texts = []
        for _, row in df.iterrows():
            row_text = " ".join(str(v) for v in row.values if pd.notna(v))
            if row_text.strip():
                texts.append(row_text.strip())

    # Return labels only if a label column was provided at CLI/env; not inferred here
    return _clean_units(texts), None

def parse_pdf(path: Path) -> List[str]:
    units = []
    if HAS_PDFPLUMBER:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                units.extend([ln.strip() for ln in txt.splitlines() if ln.strip()])
                try:
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        for row in tbl:
                            for cell in row:
                                if cell and str(cell).strip():
                                    units.append(str(cell).strip())
                except Exception:
                    pass
    elif HAS_PYPDF2:
        with path.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                txt = page.extract_text() or ""
                units.extend([ln.strip() for ln in txt.splitlines() if ln.strip()])
    else:
        raise RuntimeError("No PDF extractor available.")
    return _clean_units(units)

def parse_docx(path: Path) -> List[str]:
    units = []
    if HAS_DOCX2PY:
        with docx2python(str(path)) as doc:
            def _walk(n):
                if isinstance(n, str):
                    if n.strip():
                        units.append(n.strip())
                elif isinstance(n, list):
                    for x in n:
                        _walk(x)
            _walk(doc.body)
    elif HAS_DOCX:
        d = Document(str(path))
        for p in d.paragraphs:
            if p.text and p.text.strip():
                units.append(p.text.strip())
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        units.append(cell.text.strip())
    else:
        raise RuntimeError("No DOCX extractor available.")
    return _clean_units(units)

def parse_image(path: Path) -> List[str]:
    if not HAS_OCR:
        raise RuntimeError("pytesseract required for OCR")
    img = Image.open(path)
    txt = pytesseract.image_to_string(img) or ""
    units = [ln.strip() for ln in txt.splitlines() if ln.strip()]
    return _clean_units(units)

def parse_file(path: Path, text_col: Optional[str]) -> Tuple[List[str], Optional[List[str]]]:
    ext = path.suffix.lower()
    if ext == ".txt":
        return parse_txt(path), None
    if ext == ".csv":
        return parse_csv(path, text_col), None
    if ext == ".pdf":
        return parse_pdf(path), None
    if ext in [".docx", ".doc"]:
        return parse_docx(path), None
    if ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
        return parse_image(path), None
    raise ValueError(f"Unsupported file format: {ext}")

# --------------------------
# Data loader
# --------------------------
def load_dataset(path: str, text_col: Optional[str], label_col: Optional[str]) -> Tuple[List[str], List[str]]:
    p = Path(path).resolve()
    if not p.exists():
        for alt in [Path(path), Path.cwd() / path, Path(__file__).parent / path]:
            if alt.exists():
                p = alt.resolve()
                break
        else:
            raise FileNotFoundError(f"Dataset not found at: {path}")

    print(f"Loading dataset from: {p}")
    # If CSV and label_col provided, read labels directly
    if p.suffix.lower() == ".csv" and HAS_PANDAS:
        df = pd.read_csv(p)
        if text_col and text_col not in df.columns:
            raise ValueError(f"CSV text column '{text_col}' not found in {list(df.columns)}")
        if label_col and label_col not in df.columns:
            raise ValueError(f"CSV label column '{label_col}' not found in {list(df.columns)}")

        # Texts
        if text_col:
            texts = df[text_col].astype(str).apply(simple_normalize).tolist()
        else:
            # concatenate row if no explicit text col given
            texts = []
            for _, row in df.iterrows():
                row_text = " ".join(str(v) for v in row.values if pd.notna(v))
                texts.append(simple_normalize(row_text))

        # Labels (if provided), else heuristics
        # Labels (if provided), else heuristics
        labels = None
        if label_col and label_col in df.columns:
            labels = df[label_col].apply(normalize_label)
            # drop rows with unknown labels
            mask = labels.notna()
            dropped = (~mask).sum()
            if dropped:
                print(f"Info: dropping {dropped} rows with unmapped labels in '{label_col}'")
            texts = [t for t, keep in zip(texts, mask) if keep]
            labels = [l for l in labels[mask]]

            # final sanity check
            invalid = set(l for l in labels if l not in LABELS)
            if invalid:
                raise ValueError(f"Found invalid labels {invalid}. Allowed: {LABELS}")
        else:
            # no label column provided -> fall back to heuristic auto-labels
            labels = [heuristic_label(t) for t in texts]

        print(f"Loaded {len(texts)} samples from CSV")
        return texts, labels

    # Other formats
    texts, _ = parse_file(p, text_col)
    if not texts:
        raise ValueError("Dataset empty after parsing.")
    labels = [heuristic_label(t) for t in texts]
    print(f"Loaded {len(texts)} samples")
    return texts, labels

# --------------------------
# Rule-based features
# --------------------------
class RuleFeatureizer(BaseEstimator, TransformerMixin):
    def __init__(self, pattern_keys: Optional[List[str]] = None):
        self.pattern_keys = pattern_keys or list(PATTERNS.keys())
    def fit(self, X, y=None):
        return self
    def transform(self, X):
        import numpy as np
        feats = []
        for text in X:
            row = []
            for k in self.pattern_keys:
                p = PATTERNS[k]; matches = p.findall(text)
                row.append(int(bool(matches)))   # presence
                row.append(len(matches))         # count
            def has_any(words):
                tl = text.lower()
                return int(any(w in tl for w in words))
            row.extend([
                has_any(["credit score","income","account balance","masked pin","biometric"]),
                has_any(["agreement","supplier","customer","standing order","payment order","overdraft"]),
                has_any(["annual report","pillar 3","press","newsroom","full year results","investor"]),
                has_any(["policy","guideline","standard","sop","governance","raci","deprecated","retired"]),
            ])
            feats.append(row)
        return np.array(feats, dtype=float)

# --------------------------
# Presidio features (optional)
# --------------------------
class PresidioEntityFeaturizer(BaseEstimator, TransformerMixin):
    def __init__(self, language="en", top_k_entities=16):
        self.language = language
        self.top_k_entities = top_k_entities
        self.entity_order_ = []
    def fit(self, X, y=None):
        if not (USE_PRESIDIO and HAS_PRESIDIO):
            return self
        from collections import Counter
        c = Counter()
        for txt in X:
            try:
                results = _ANALYZER.analyze(text=txt, language=self.language)
                c.update(r.entity_type for r in results)
            except Exception:
                pass
        self.entity_order_ = [et for et, _ in c.most_common(self.top_k_entities)]
        return self
    def transform(self, X):
        import numpy as np
        if not (USE_PRESIDIO and HAS_PRESIDIO) or not self.entity_order_:
            return np.zeros((len(X), len(self.entity_order_)*2+1))
        feats = []
        for txt in X:
            try:
                results = _ANALYZER.analyze(text=txt, language=self.language)
            except Exception:
                results = []
            types = [r.entity_type for r in results]
            row = []
            # counts per entity type
            for et in self.entity_order_:
                row.append(float(types.count(et)))
            # presence per entity type
            for et in self.entity_order_:
                row.append(1.0 if et in types else 0.0)
            # total entities
            row.append(float(len(types)))
            feats.append(row)
        return np.array(feats, dtype=float)

# --------------------------
# Feature builders
# --------------------------
def build_features(use_hash=False) -> FeatureUnion:
    rules = RuleFeatureizer()
    if use_hash:
        bow = HashingVectorizer(
            n_features=2**16, alternate_sign=False,
            lowercase=True, token_pattern=r"(?u)\b[\w'-]+\b",
            strip_accents='unicode'
        )
        text_block = ("bow", bow)
    else:
        tfidf = TfidfVectorizer(
            ngram_range=(1, NGRAM_MAX), min_df=MIN_DF, max_df=MAX_DF,
            lowercase=True, token_pattern=r"(?u)\b[\w'-]+\b",
            strip_accents='unicode', stop_words=None
        )
        text_block = ("tfidf", tfidf)

    blocks = [text_block, ("rules", rules)]
    if USE_PRESIDIO and HAS_PRESIDIO:
        blocks.append(("presidio", PresidioEntityFeaturizer(language="en", top_k_entities=16)))
    return FeatureUnion(blocks)

# --------------------------
# Classifier chooser
# --------------------------
def build_classifier(name: str):
    if name == "logreg":
        # Optionally set class_weight="balanced" if your dataset is skewed
        return LogisticRegression(max_iter=300, solver="lbfgs", multi_class="auto", random_state=RANDOM_STATE)
    if name == "rf":
        return RandomForestClassifier(n_estimators=300, random_state=RANDOM_STATE, class_weight="balanced")
    raise ValueError(f"Unknown MODEL: {name}")

# --------------------------
# Train / Eval
# --------------------------
@dataclass
class TrainResult:
    report: str
    conf_mat: List[List[int]]
    model_path: str

def train_and_eval(data_path: str, text_col: Optional[str], label_col: Optional[str]) -> TrainResult:
    texts, labels = load_dataset(data_path, text_col, label_col)

    # Stratified split; fall back if class counts too small
    try:
        X_train, X_test, y_train, y_test = train_test_split(
            texts, labels, test_size=TEST_SIZE, random_state=RANDOM_STATE, stratify=labels
        )
    except ValueError:
        X_train, X_test, y_train, y_test = train_test_split(
            texts, labels, test_size=TEST_SIZE, random_state=RANDOM_STATE, stratify=None
        )

    clf = build_classifier(MODEL_CHOICE)

    use_hash = False
    try:
        pipe = Pipeline([("features", build_features(use_hash)), ("clf", clf)])
        pipe.fit(X_train, y_train)
    except ValueError as e:
        msg = str(e).lower()
        if "empty vocabulary" in msg or "no terms remain" in msg:
            use_hash = True
            pipe = Pipeline([("features", build_features(use_hash)), ("clf", clf)])
            pipe.fit(X_train, y_train)
        else:
            raise

    y_pred = pipe.predict(X_test)
    report = classification_report(y_test, y_pred, digits=3, labels=LABELS, zero_division=0)
    cm = confusion_matrix(y_test, y_pred, labels=LABELS).tolist()

    out_path = str((Path(__file__).parent / "sensitivity_classifier.joblib").resolve())
    joblib.dump({"pipeline": pipe, "labels": LABELS}, out_path)

    # Also save metrics to a json for convenience
    metrics_path = Path(__file__).parent / "model_metrics.json"
    with open(metrics_path, "w", encoding="utf-8") as f:
        json.dump({"report": report, "confusion_matrix": cm}, f, indent=2)

    return TrainResult(report=report, conf_mat=cm, model_path=out_path)

# --------------------------
# CLI
# --------------------------
def parse_args(argv: List[str]):
    data_arg = None
    text_col = CSV_TEXT_COLUMN
    label_col = CSV_LABEL_COLUMN
    if "--data" in argv:
        i = argv.index("--data")
        if i + 1 < len(argv):
            data_arg = argv[i + 1]
    if "--text-col" in argv:
        i = argv.index("--text-col")
        if i + 1 < len(argv):
            text_col = argv[i + 1]
    if "--label-col" in argv:
        i = argv.index("--label-col")
        if i + 1 < len(argv):
            label_col = argv[i + 1]
    return data_arg, text_col, label_col

def main(argv: List[str] = None):
    argv = argv or sys.argv[1:]
    data_arg, text_col, label_col = parse_args(argv)
    data_path = data_arg or DATA_PATH

    if "--train" in argv or len(argv) == 0:
        res = train_and_eval(data_path, text_col, label_col)
        print("\n=== Sensitivity Classifier (Hybrid: rules + TF-IDF + ML) ===")
        print(f"Model: {MODEL_CHOICE}")
        print(f"Data:  {data_path}")
        if text_col:  print(f"Text col:   {text_col}")
        if label_col: print(f"Label col:  {label_col}")
        print("\n--- Classification Report ---")
        print(res.report)
        print("\n--- Confusion Matrix [rows=true, cols=pred] (order C1,C2,C3,C4) ---")
        print(json.dumps(res.conf_mat))
        print(f"\nModel saved to: {res.model_path}")
    else:
        print("Usage:")
        print("  python ml_setup.py --train --data /path/to/data.csv --text-col text --label-col label")
        print("  CSV_TEXT_COLUMN=text CSV_LABEL_COLUMN=label python ml_setup.py --train --data data.csv")
        print("\nSupported formats (for ad-hoc use): TXT, CSV, PDF, DOCX, PNG, JPG, JPEG")
        print("If no label column is given, labels are assigned by heuristics (C1–C4).")

if __name__ == "__main__":
    main()