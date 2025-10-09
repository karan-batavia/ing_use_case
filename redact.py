#!/usr/bin/env python3
import re
import os
import json
import csv
from typing import Dict, List, Tuple, Any, Optional
from dataclasses import dataclass
from pathlib import Path

# ============================
#  Config (env toggles)
# ============================
ONLY_USER_PATTERNS = os.environ.get("ONLY_USER_PATTERNS", "1") == "1"
USE_SPACY = os.environ.get("USE_SPACY", "1") == "1"
SPACY_MODEL = os.environ.get("SPACY_MODEL", "en_core_web_sm")
USE_PERSON_REGEX_FALLBACK = os.environ.get("USE_PERSON_REGEX_FALLBACK", "1") == "1"

# Conservative name regex for fallback (does NOT change your own patterns)
PERSON_FALLBACK_REGEX = (
    r"\b(?:[A-ZÀ-ÖØ-Ý][a-zà-öø-ÿ]+(?:[-'][A-ZÀ-ÖØ-Ý][a-zà-öø-ÿ]+)?\s+){1,3}"
    r"[A-ZÀ-ÖØ-Ý][a-zà-ö-ÿ]+(?:[-'][A-ZÀ-ÖØ-Ý][a-zà-öø-ÿ]+)?\b"
)

# ============================
#  Optional deps (parsers)
# ============================
HAS_PANDAS = False
try:
    import pandas as pd
    HAS_PANDAS = True
except Exception:
    pass

HAS_PDFPLUMBER = False
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except Exception:
    pass

HAS_PYPDF2 = False
try:
    import PyPDF2
    HAS_PYPDF2 = True
except Exception:
    pass

HAS_DOCX2PY = False
try:
    from docx2python import docx2python
    HAS_DOCX2PY = True
except Exception:
    pass

HAS_DOCX = False
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    pass

HAS_OCR = False
try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except Exception:
    pass

# ============================
#  Load your regex patterns
# ============================
USER_REGEX_OK = False
user_regex = None
try:
    from preprocessing import regex_queries as user_regex
    if isinstance(getattr(user_regex, "PATTERNS", None), dict):
        USER_REGEX_OK = True
        print(f"[INFO] Using user regex patterns from preprocessing.regex_queries (count={len(user_regex.PATTERNS)})")
    else:
        print("[WARN] preprocessing.regex_queries found but PATTERNS missing or not a dict.")
except Exception as _e:
    print(f"[WARN] Could not import preprocessing.regex_queries: {_e}")

# ============================
#  Optional spaCy NER
# ============================
HAS_SPACY = False
_SPACY_NLP = None
if USE_SPACY:
    try:
        import spacy
        HAS_SPACY = True
    except Exception as _e:
        print(f"[WARN] spaCy not available ({_e}). Proceeding without NER.")
        HAS_SPACY = False

def _get_spacy():
    """Lazy-load spaCy model once."""
    global _SPACY_NLP
    if not (USE_SPACY and HAS_SPACY):
        return None
    if _SPACY_NLP is None:
        try:
            _SPACY_NLP = spacy.load(SPACY_MODEL, disable=["tagger","parser","lemmatizer","textcat"])
            print(f"[INFO] spaCy NER enabled with model: {SPACY_MODEL}")
        except Exception as _e:
            print(f"[WARN] Failed to load spaCy model '{SPACY_MODEL}': {_e}. Disabling NER.")
            _SPACY_NLP = None
    return _SPACY_NLP

# ============================
#  Fallback regex patterns
# ============================
FALLBACK_PATTERNS: Dict[str, str] = {
    "EMAIL": r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
    "PHONE_EU": r"(?:\+\d{1,3}\s?)?(?:\d[\s-]?){9,}",
    "SSN_LIKE": r"\b\d{2}[.\-/]\d{2}[.\-/]\d{2}[- ]?\d{2,4}[.\-]?\d{0,2}\b|\b\d{6}[- ]?\d{2,4}[.\-]?\d{0,2}\b",
    "IBAN": r"\b[A-Z]{2}\d{2}[ ]?(?:[A-Z0-9]{3,4}[ ]?){3,7}\b",
    "ACCOUNT_NUM": r"\b(?:acct|account)\s*\d{3,}\b",
    "AMOUNT": r"(?:USD|EUR|GBP|€|\$|£)\s?\d{1,3}(?:[, \u00A0]\d{3})*(?:\.\d{2})?",
    "DOB": r"\b\d{4}-\d{2}-\d{2}\b",
    "NATIONAL_ID": r"\b(?:national (?:id|number)|id)[:\s-]*[A-Z0-9.\-]{6,}\b",
    "BIOMETRIC": r"\b(FaceID|fingerprint|iris|biometric)\b",
    "IBAN_PLAIN": r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b",
    # PERSON is added by spaCy (and optional fallback); we won't touch your PATTERNS.
}

# Normalize common alias keys to PERSON if present in user patterns
PERSON_ALIASES = {"NAME", "PERSON_NAME", "FULL_NAME"}

def _compile_patterns() -> Dict[str, re.Pattern]:
    if USER_REGEX_OK and ONLY_USER_PATTERNS:
        base = dict(user_regex.PATTERNS)  # yours only
        print("[INFO] ONLY_USER_PATTERNS=1 → using only user patterns.")
    elif USER_REGEX_OK:
        base = {**FALLBACK_PATTERNS, **user_regex.PATTERNS}  # merge
        print("[INFO] ONLY_USER_PATTERNS=0 → merging user patterns over fallbacks.")
    else:
        base = dict(FALLBACK_PATTERNS)
        print("[INFO] Using FALLBACK_PATTERNS (no user patterns found).")

    # Alias → PERSON
    for alias in list(base.keys()):
        if alias.upper() in PERSON_ALIASES and "PERSON" not in base:
            base["PERSON"] = base[alias]

    return {k: re.compile(v, re.IGNORECASE) for k, v in base.items()}

# ============================
#  Explanations & priorities
# ============================
EXPLANATIONS: Dict[str, str] = {
    "EMAIL": "Email addresses are personally identifiable information (PII).",
    "PHONE_EU": "Phone numbers are PII and can identify or contact an individual.",
    "SSN_LIKE": "National/SSN-like identifiers are highly sensitive government-issued IDs.",
    "IBAN": "IBAN corresponds to a bank account identifier and is financial PII.",
    "IBAN_PLAIN": "IBAN corresponds to a bank account identifier and is financial PII.",
    "ACCOUNT_NUM": "Account numbers relate to financial PII and should be masked.",
    "AMOUNT": "Monetary amounts may be business-sensitive when linked to people or transactions.",
    "DOB": "Date of birth is personal data that can identify an individual.",
    "NATIONAL_ID": "National ID/number is a highly sensitive personal identifier.",
    "BIOMETRIC": "Biometric references relate to sensitive authentication factors.",
    "PERSON": "Person names are personal data (PII).",
}

TYPE_PRIORITY = [
    "SSN_LIKE", "NATIONAL_ID", "IBAN", "IBAN_PLAIN",
    "ACCOUNT_NUM", "PERSON", "EMAIL", "PHONE_EU", "DOB", "AMOUNT", "BIOMETRIC"
]

# ============================
#  Core classes
# ============================
@dataclass
class MatchItem:
    type: str
    start: int
    end: int
    text: str

class Redactor:
    def __init__(self, patterns: Optional[Dict[str, re.Pattern]] = None, debug: bool = False):
        self.patterns = patterns or _compile_patterns()
        self.counters: Dict[str, int] = {}
        self.debug = debug
        self._person_fallback_re = re.compile(PERSON_FALLBACK_REGEX) if USE_PERSON_REGEX_FALLBACK else None

    def _placeholder(self, typ: str) -> str:
        n = self.counters.get(typ, 0) + 1
        self.counters[typ] = n
        return f"[{typ}_{n:03d}]"

    def _find_person_spacy(self, text: str) -> List[MatchItem]:
        nlp = _get_spacy()
        if nlp is None:
            return []
        try:
            doc = nlp(text)
        except Exception:
            return []
        out: List[MatchItem] = []
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                out.append(MatchItem(type="PERSON", start=ent.start_char, end=ent.end_char, text=ent.text))
        return out

    def _find_person_fallback(self, text: str) -> List[MatchItem]:
        if not self._person_fallback_re:
            return []
        return [MatchItem("PERSON", m.start(), m.end(), m.group(0))
                for m in self._person_fallback_re.finditer(text)]

    def find_entities(self, text: str) -> List[MatchItem]:
        found: List[MatchItem] = []

        # 1) Regex matches (yours + fallbacks)
        for typ, pat in self.patterns.items():
            for m in pat.finditer(text):
                found.append(MatchItem(type=typ, start=m.start(), end=m.end(), text=m.group(0)))

        # 2) spaCy PERSON
        found.extend(self._find_person_spacy(text))

        # 3) PERSON fallback if none found
        if USE_PERSON_REGEX_FALLBACK and not any(mi.type == "PERSON" for mi in found):
            found.extend(self._find_person_fallback(text))

        if self.debug:
            print("---- RAW MATCHES ----")
            for mi in found:
                print(f"{mi.type:<10} [{mi.start:>4},{mi.end:<4}]  {mi.text!r}")

        if not found:
            return []

        # Sort: longer first, then higher priority, then earlier start
        def prio(mi: MatchItem) -> Tuple[int, int, int]:
            length = mi.end - mi.start
            pidx = TYPE_PRIORITY.index(mi.type) if mi.type in TYPE_PRIORITY else len(TYPE_PRIORITY)
            return (-length, pidx, mi.start)
        found.sort(key=prio)

        # Greedy non-overlapping selection
        selected: List[MatchItem] = []
        occupied: List[Tuple[int, int]] = []
        for mi in found:
            overlaps = any(not (mi.end <= s or mi.start >= e) for (s, e) in occupied)
            if not overlaps:
                selected.append(mi)
                occupied.append((mi.start, mi.end))

        selected.sort(key=lambda x: x.start)

        if self.debug:
            print("---- SELECTED ----")
            for mi in selected:
                print(f"{mi.type:<10} [{mi.start:>4},{mi.end:<4}]  {mi.text!r}")

        return selected

    def redact_text(self, text: str) -> Dict[str, Any]:
        entities = self.find_entities(text)
        if not entities:
            return {"text": text, "mapping": [], "index": {}}

        out_chars: List[str] = []
        last = 0
        mapping: List[Dict[str, Any]] = []
        index: Dict[str, str] = {}

        for mi in entities:
            out_chars.append(text[last:mi.start])
            ph = self._placeholder(mi.type)
            out_chars.append(ph)
            mapping.append({
                "type": mi.type,
                "placeholder": ph,
                "original": mi.text,
                "span": [mi.start, mi.end],
                "explanation": EXPLANATIONS.get(mi.type, "Redacted sensitive entity."),
            })
            index[ph] = mi.text
            last = mi.end

        out_chars.append(text[last:])
        redacted = "".join(out_chars)
        return {"text": redacted, "mapping": mapping, "index": index}

    def redact_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        return [self.redact_text(ln) for ln in lines]

# ============================
#  Multi-format readers
# ============================
def _clean_units(units: List[str]) -> List[str]:
    out = []
    for u in units:
        if not u:
            continue
        s = str(u).strip()
        if not s:
            continue
        # drop very short tokens that aren't likely entities and have no digits/symbols
        if len(s) < 3 and not re.search(r"[0-9@€$£]", s):
            continue
        out.append(s)
    return out

def parse_txt(path: Path) -> List[str]:
    return _clean_units([ln.rstrip("\n") for ln in path.read_text(encoding="utf-8", errors="ignore").splitlines()])

def parse_csv(path: Path, text_col: Optional[str]) -> List[str]:
    if not HAS_PANDAS:
        raise RuntimeError("pandas is required to read CSV files.")
    df = pd.read_csv(path)
    if text_col:
        if text_col not in df.columns:
            raise ValueError(f"CSV text column '{text_col}' not found in {list(df.columns)}")
        texts = df[text_col].astype(str).tolist()
    else:
        # concat row values into one string
        texts = []
        for _, row in df.iterrows():
            val = " ".join(str(v) for v in row.values if pd.notna(v))
            texts.append(val)
    return _clean_units(texts)

def parse_pdf(path: Path) -> List[str]:
    units: List[str] = []
    if HAS_PDFPLUMBER:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt:
                    units.extend([ln.strip() for ln in txt.splitlines() if ln.strip()])
                # also try tables
                try:
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        for row in tbl:
                            for cell in row:
                                if cell and str(cell).strip():
                                    units.append(str(cell).strip())
                except Exception:
                    pass
    elif HAS_PYPDF2:
        with path.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                txt = page.extract_text() or ""
                if txt:
                    units.extend([ln.strip() for ln in txt.splitlines() if ln.strip()])
    else:
        raise RuntimeError("No PDF extractor available. Install pdfplumber or PyPDF2.")
    return _clean_units(units)

def parse_docx(path: Path) -> List[str]:
    units: List[str] = []
    if HAS_DOCX2PY:
        with docx2python(str(path)) as doc:
            def _walk(n):
                if isinstance(n, str):
                    if n.strip(): units.append(n.strip())
                elif isinstance(n, list):
                    for x in n: _walk(x)
            _walk(doc.body)
    elif HAS_DOCX:
        d = Document(str(path))
        for p in d.paragraphs:
            if p.text and p.text.strip():
                units.append(p.text.strip())
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        units.append(cell.text.strip())
    else:
        raise RuntimeError("No DOCX extractor available. Install python-docx or docx2python.")
    return _clean_units(units)

def parse_image(path: Path) -> List[str]:
    if not HAS_OCR:
        raise RuntimeError("OCR requires pillow and pytesseract (and Tesseract installed).")
    img = Image.open(path)
    txt = pytesseract.image_to_string(img) or ""
    units = [ln.strip() for ln in txt.splitlines() if ln.strip()]
    return _clean_units(units)

def parse_file(path: Path, text_col: Optional[str]) -> List[str]:
    ext = path.suffix.lower()
    if ext == ".txt":
        return parse_txt(path)
    if ext == ".csv":
        return parse_csv(path, text_col)
    if ext == ".pdf":
        return parse_pdf(path)
    if ext in [".docx", ".doc"]:
        return parse_docx(path)
    if ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
        return parse_image(path)
    # Fallback: try reading as text
    try:
        return parse_txt(path)
    except Exception:
        raise ValueError(f"Unsupported file format: {ext}")

# ============================
#  Output helpers
# ============================
def _ensure_parent(path: Optional[str]):
    if not path:
        return
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)

def _write_outputs(results, out_path, map_path, csv_path):
    _ensure_parent(out_path)
    _ensure_parent(map_path)
    _ensure_parent(csv_path)

    if out_path:
        with open(out_path, "w", encoding="utf-8") as f:
            for r in results:
                f.write(r["text"] + "\n")

    if map_path:
        with open(map_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)

    if csv_path:
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["line_idx","type","placeholder","original","span_start","span_end","explanation"])
            for i, r in enumerate(results):
                if not r["mapping"]:
                    w.writerow([i, "", "", "", "", "", "", ""])
                else:
                    for m in r["mapping"]:
                        w.writerow([i, m["type"], m["placeholder"], m["original"], m["span"][0], m["span"][1], m["explanation"]])

# ============================
#  CLI
# ============================
def main():
    import argparse
    ap = argparse.ArgumentParser(description="Redact entities from TXT/CSV/PDF/DOCX/Images (regex + spaCy PERSON + fallback).")
    ap.add_argument("--in", dest="infile", required=True, help="Input file (.txt, .csv, .pdf, .docx/.doc, image)")
    ap.add_argument("--out", dest="outfile", required=True, help="Output redacted text file")
    ap.add_argument("--map", dest="mapfile", default=None, help="JSON with redaction mappings")
    ap.add_argument("--csv", dest="csvfile", default=None, help="CSV with flattened mappings")
    ap.add_argument("--text-col", dest="text_col", default=None, help="CSV column name to read as text (if omitted, concatenates row)")
    ap.add_argument("--limit", type=int, default=0, help="Optional: process only first N records/lines")
    ap.add_argument("--debug", action="store_true", help="Enable verbose output (prints found entities)")
    args = ap.parse_args()

    p = Path(args.infile).expanduser().resolve()
    if not p.exists():
        ap.error(f"Input file not found: {p}")

    try:
        lines = parse_file(p, args.text_col)
    except Exception as e:
        ap.error(f"Failed to parse input '{p}': {e}")

    if args.limit and args.limit > 0:
        lines = lines[:args.limit]

    red = Redactor(debug=args.debug)
    results = red.redact_lines(lines)
    _write_outputs(results, args.outfile, args.mapfile, args.csvfile)

if __name__ == "__main__":
    main()